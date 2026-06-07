"""Generate a DOCX version of the CV mimicking the HTML layout.

Usage (examples):
  python docx.py                         # uses src/en/language.json
  python docx.py --lang de               # uses src/de/language.json
  python docx.py --file src/de/language1.json --out output/cv-de1.docx
  python docx.py --assets test/CV_files  # custom assets folder for images

Install dependency:  pip install python-docx
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

# Avoid shadowing the external python-docx package with this script's filename
sys.path = [p for p in sys.path if p not in ("", os.getcwd())]

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ModuleNotFoundError as exc:
    raise SystemExit(
        "Missing dependency 'python-docx'. Install it with: pip install python-docx"
    ) from exc


def load_data(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


ADDRESS = "Karlsbader Str. 7, Ludwigshafen am Rhein, 67065"
PHONE = "+49 176 70575532"
EMAIL = "alejandropriv@outlook.com"
LINKEDIN = "https://www.linkedin.com/in/alejandroprietov/"
GITHUB = "https://github.com/alejandropriv"
NAME = "Alejandro Prieto"
PHOTO_FILE = "bewerbung01.jpg"
SIGNATURE_FILE = "firma.jpg"


def add_heading(target, text: str, level: int = 1):
    if not text:
        return
    if hasattr(target, "add_heading"):
        target.add_heading(text, level=level)
    else:
        target.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(str(item))


def render_skills(doc: Document, data: dict):
    sk = data.get("SK100", {})
    add_heading(doc, sk.get("title", "Skills"), level=2)
    for name, info in sk.get("training", {}).items():
        if info.get("enable"):
            stars = "★" * int(info.get("value", 0))
            p = doc.add_paragraph(f"{name} {stars}")
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(8)


def render_training(doc: Document, data: dict):
    ct = data.get("CT100", {})
    add_heading(doc, ct.get("title", "Training"), level=2)
    for item in ct.get("certifications", {}).keys():
        p = doc.add_paragraph(str(item))
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)

    cc = data.get("CC100", {})
    for item in cc.get("certifications", {}).keys():
        p = doc.add_paragraph(str(item))
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)


def render_languages(doc: Document, data: dict):
    lang = data.get("SKL100", {})
    add_heading(doc, lang.get("title", "Languages"), level=2)
    for name, value in lang.get("languages", {}).items():
        p = doc.add_paragraph(f"{name} {'★'*int(value)}")
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)


def render_header(doc: Document, data: dict):
    add_heading(doc, NAME, level=1)
    if doc.paragraphs:
        p = doc.paragraphs[-1]
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(14)
    if cp := data.get("contactPresentation"):
        p = doc.add_paragraph(cp)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(14)



def render_work_experience(doc: Document, data: dict):
    add_heading(doc, data.get("WETitleGeneral", "Work Experience"), level=2)
    for item in data.get("WE", []):
        if not item.get("WEEnable"):
            continue
        if item.get("WETitle") == "divider":
            continue
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(item.get("WETitle", ""))
        title_run.bold = True

        date_p = doc.add_paragraph(item.get("WETime", ""))
        if date_p.runs:
            date_run = date_p.runs[0]
            date_run.italic = True
            date_run.font.name = "Arial"
            date_run.font.size = Pt(10)

        for desc in item.get("WEDescription", []):
            if desc.get("enabled"):
                prefix = f"{desc.get('title')}: " if desc.get("title") else ""
                p = doc.add_paragraph(f"{prefix}{desc.get('text', '')}", style="List Bullet")
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)


def render_contracts(doc: Document, data: dict):
    c_title = data.get("CTitleGeneral", {}).get("CTitle")
    if not c_title or data.get("CTitleGeneral", {}).get("CTEnable") != "true":
        return
    add_heading(doc, c_title, level=2)
    for item in data.get("C", []):
        if item.get("CEnable") != "true":
            continue
        if item.get("CTitle") == "divider":
            continue
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(item.get("CTitle", ""))
        title_run.bold = True

        date_p = doc.add_paragraph(item.get("CTime", ""))
        if date_p.runs:
            date_run = date_p.runs[0]
            date_run.italic = True
            date_run.font.name = "Arial"
            date_run.font.size = Pt(10)

        for desc in item.get("CDescription", []):
            if desc.get("enabled"):
                prefix = f"{desc.get('title')}: " if desc.get("title") else ""
                p = doc.add_paragraph(f"{prefix}{desc.get('text', '')}", style="List Bullet")
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)


def render_studies(doc: Document, data: dict):
    add_heading(doc, data.get("STitleGeneral", "Studies"), level=2)
    for item in data.get("S", []):
        if not item.get("SEnable"):
            continue
        if item.get("STitle") == "divider":
            continue
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(item.get("STitle", ""))
        title_run.bold = True

        date_p = doc.add_paragraph(item.get("STime", ""))
        if date_p.runs:
            date_run = date_p.runs[0]
            date_run.italic = True
            date_run.font.name = "Arial"
            date_run.font.size = Pt(10)

        for desc in item.get("SDescription", []):
            if desc.get("enabled"):
                prefix = f"{desc.get('title')}: " if desc.get("title") else ""
                p = doc.add_paragraph(f"{prefix}{desc.get('text', '')}", style="List Bullet")
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)


def find_asset(filename: str, assets_dir: Path | None) -> Path | None:
    candidates = []
    if assets_dir:
        candidates.append(assets_dir / filename)
    candidates += [
        Path("src/images") / filename,
        Path("test/CV_files") / filename,
        Path("CV_files") / filename,
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def set_cell_border(cell, color: str = "FFFFFF"):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_tag = qn(f"w:{edge}")
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def build_doc(data: dict, assets_dir: Path | None) -> Document:
    doc = Document()

    # Page setup: reduced margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Base styles: Arial 10pt
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)

    for h in ["Heading 1", "Heading 2", "Heading 3"]:
        if h in doc.styles:
            s = doc.styles[h]
            s.font.name = "Arial"
            s.font.size = Pt(12)
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.font.bold = True

    # Two-column layout using a full-width table with white borders
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    col_left, col_right = table.columns[0], table.columns[1]
    col_left.width = int(page_width * 0.25)
    col_right.width = int(page_width * 0.75)

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF")

    # Left column: photo, contact, skills, training, languages
    photo_path = find_asset(PHOTO_FILE, assets_dir)
    if photo_path:
        try:
            left.add_paragraph().add_run().add_picture(str(photo_path), width=Inches(1.6))
        except Exception:
            pass

    for text in [ADDRESS, PHONE, EMAIL, LINKEDIN, GITHUB]:
        p = left.add_paragraph(text)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(8)

    render_skills(left, data)
    render_training(left, data)
    render_languages(left, data)

    # Right column: name header, summary, experiences, contracts, studies, signature
    render_header(right, data)
    render_work_experience(right, data)
    render_contracts(right, data)
    render_studies(right, data)

    sig_path = find_asset(SIGNATURE_FILE, assets_dir)
    if sig_path:
        try:
            right.add_paragraph().add_run().add_picture(str(sig_path), width=Inches(2.5))
        except Exception:
            pass

    return doc


def main():
    parser = argparse.ArgumentParser(description="Generate CV DOCX from JSON data")
    parser.add_argument("--lang", default="en", help="Language folder (en, de, de1, es, en1)")
    parser.add_argument("--file", type=Path, help="Override JSON file path")
    parser.add_argument("--out", type=Path, default=Path("output/alejandro-prieto-cv.docx"), help="Output DOCX path")
    parser.add_argument("--assets", type=Path, help="Folder containing images (photo, signature)")
    args = parser.parse_args()

    data_file = args.file or Path(f"src/{args.lang}/language.json")

    if not data_file.exists():
        raise SystemExit(f"Language file not found: {data_file}")

    data = load_data(data_file)

    args.out.parent.mkdir(parents=True, exist_ok=True)

    doc = build_doc(data, args.assets)
    doc.save(args.out)

    print(f"DOCX generated: {args.out}")


if __name__ == "__main__":
    main()
